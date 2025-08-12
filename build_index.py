import json
import faiss
import yaml
import numpy as np
from pathlib import Path
from sentence_transformers import SentenceTransformer
from pypdf import PdfReader
import docx
import re
import warnings
from pypdf.errors import PdfReadWarning
import pypdfium2 as pdfium

warnings.filterwarnings("ignore", category=PdfReadWarning)

# Load config
with open("config.yaml", "r", encoding="utf-8") as f:
    config = yaml.safe_load(f)

DOCS_DIR = Path(config["docs_dir"])
INDEX_DIR = DOCS_DIR / config["index_dir"]
INDEX_DIR.mkdir(exist_ok=True)
MODEL_NAME = config["model"]["name"]
DEVICE = config["model"]["device"]
CHUNK_SIZE = config["chunking"]["size"]
CHUNK_OVERLAP = config["chunking"]["overlap"]
FILE_TYPES = config["file_types"]


def pdf_extract_text(path: Path) -> str:
    # Try pypdf
    try:
        reader = PdfReader(str(path))
        parts = []
        for p in reader.pages:
            t = (p.extract_text() or "").strip()
            parts.append(t)
        text = "\n".join(parts).strip()
        if len(text) >= 40:  # good enough → use it
            return text
    except Exception:
        pass

    # Fallback: pypdfium2 text extraction
    try:
        doc = pdfium.PdfDocument(str(path))
        out = []
        for i in range(len(doc)):
            page = doc[i]
            textpage = page.get_textpage()
            out.append(textpage.get_text_range())
        return "\n".join(out).strip()
    except Exception:
        return ""

# Helper: text extraction
def extract_text(file_path: Path) -> str:
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return pdf_extract_text(file_path)
    elif ext == ".docx":
        doc = docx.Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)
    elif ext in (".txt", ".md"):
        return file_path.read_text(encoding="utf-8", errors="ignore")
    return ""

# Helper: chunking
def chunk_text(text: str):
    words = re.split(r"\s+", text)
    chunks = []
    for i in range(0, len(words), CHUNK_SIZE - CHUNK_OVERLAP):
        chunk = " ".join(words[i:i + CHUNK_SIZE])
        if chunk.strip():
            chunks.append(chunk)
    return chunks


# Build index
model = SentenceTransformer(MODEL_NAME, device=DEVICE)
docs_meta = []
vectors = []

for file_path in DOCS_DIR.rglob("*"):
    if file_path.suffix.lower() in FILE_TYPES:
        text = extract_text(file_path)
        if not text.strip():
            continue
        for chunk in chunk_text(text):
            emb = model.encode(chunk)
            vectors.append(emb)
            docs_meta.append({
                "source": str(file_path),
                "content": chunk
            })

if not vectors:
    raise ValueError("No documents found to index.")

# Stack to 2D float32
vecs = np.vstack(vectors).astype("float32")

# Single normalization pass (for cosine via IndexFlatIP)
vecs /= (np.linalg.norm(vecs, axis=1, keepdims=True) + 1e-12)

dim = vecs.shape[1]
index = faiss.IndexFlatIP(dim)  # cosine similarity when vectors are normalized
index.add(vecs)

faiss.write_index(index, str(INDEX_DIR / "faiss_index.bin"))
with open(INDEX_DIR / "metadata.json", "w", encoding="utf-8") as f:
    json.dump(docs_meta, f, ensure_ascii=False, indent=2)

print(f"Indexed {len(docs_meta)} chunks from {DOCS_DIR}")
